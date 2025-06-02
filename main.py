import streamlit as st
import os
import json
import pandas as pd
import re
import PyPDF2
import docx
import tempfile
from datetime import datetime
from io import BytesIO
import numpy as np
import anthropic
from dotenv import load_dotenv
import time
import hashlib
import hmac
import traceback
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def load_environment_variables():
    """Load environment variables with proper error handling"""
    try:
        load_dotenv()
    except Exception as e:
        logger.warning(f"Could not load .env file: {e}")

# Load environment variables
load_environment_variables()

def safe_get_secret(key, default_value):
    """Safely get secrets without crashing if secrets.toml doesn't exist"""
    try:
        return st.secrets.get(key, default_value)
    except:
        return default_value

# Define valid users - using environment variables with fallbacks
VALID_USERS = {
    "smartworks_admin": os.getenv("SMARTWORKS_ADMIN_PASSWORD") or safe_get_secret("SMARTWORKS_ADMIN_PASSWORD", "sw2025!"),
    "client_manager": os.getenv("CLIENT_MANAGER_PASSWORD") or safe_get_secret("CLIENT_MANAGER_PASSWORD", "cm2024!"),
    "operations": os.getenv("OPERATIONS_PASSWORD") or safe_get_secret("OPERATIONS_PASSWORD", "ops2024!"),
    "ansh.arora1@sworks.co.in": os.getenv("ANSH_PASSWORD") or safe_get_secret("ANSH_PASSWORD", "ansh1529")
}

# Initialize Claude API client
@st.cache_resource
def init_claude_client():
    """Initialize Claude client with proper error handling"""
    try:
        api_key = os.getenv("ANTHROPIC_API_KEY") or safe_get_secret("ANTHROPIC_API_KEY", None)
        if not api_key:
            st.error("Please set your ANTHROPIC_API_KEY in the environment variables or Streamlit secrets")
            return None
        return anthropic.Anthropic(api_key=api_key)
    except Exception as e:
        st.error(f"Failed to initialize Claude client: {e}")
        return None

# Initialize session state
def init_session_state():
    """Initialize session state variables with authentication"""
    default_values = {
        'successful_resumes': [],
        'candidate_df': None,
        'shortlisted_df': None,
        'job_requirements': None,
        'logs': [],
        'current_job_title': "job_position",
        'total_files': 0,
        'successful_count': 0,
        'failed_count': 0,
        'top_candidates': [],
        'weights': {
            "experience": 0.30,
            "skills": 0.40,
            "education": 0.20,
            "certification": 0.10
        },
        'password_correct': False,
        'authenticated_user': None,
        'processing_complete': False,
        'current_job_description': None,
        'processing_seed': None
    }
    
    for key, value in default_values.items():
        if key not in st.session_state:
            st.session_state[key] = value

def generate_processing_seed(job_description, weights):
    """Generate a consistent seed based on job description and weights for reproducible results"""
    try:
        # Create a consistent string from job description and weights
        seed_string = f"{job_description.strip()}"
        seed_string += f"_exp:{weights['experience']}_skills:{weights['skills']}_edu:{weights['education']}_cert:{weights['certification']}"
        
        # Generate a hash to create a numeric seed
        seed_hash = hashlib.md5(seed_string.encode()).hexdigest()
        # Convert first 8 characters of hash to integer
        seed = int(seed_hash[:8], 16) % 1000000
        return seed
    except Exception as e:
        logger.warning(f"Could not generate seed: {e}")
        return 42  # Default seed
    
# Text extraction functions with better error handling
def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file with enhanced error handling"""
    text = ""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        logger.info(f"📄 PDF has {len(pdf_reader.pages)} pages")
        
        for page_num in range(len(pdf_reader.pages)):
            try:
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                    logger.info(f"✅ Extracted text from page {page_num + 1}: {len(page_text)} characters")
                else:
                    logger.warning(f"⚠️ No text found on page {page_num + 1}")
            except Exception as page_error:
                logger.error(f"❌ Error extracting page {page_num + 1}: {str(page_error)}")
                continue
                
        if not text.strip():
            raise Exception("No text could be extracted from any page")
            
        logger.info(f"✅ Total extracted text: {len(text)} characters")
        return text
        
    except Exception as e:
        error_msg = f"PDF extraction failed: {str(e)}"
        logger.error(f"❌ {error_msg}")
        raise Exception(error_msg)

def extract_text_from_docx(docx_file):
    """Extract text from DOCX file with enhanced error handling"""
    text = ""
    try:
        doc = docx.Document(docx_file)
        
        # Extract from paragraphs
        paragraph_count = 0
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
                paragraph_count += 1
        
        logger.info(f"📝 Extracted text from {paragraph_count} paragraphs")
        
        # Extract from tables
        table_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
            table_count += 1
            
        logger.info(f"📊 Extracted text from {table_count} tables")
        
        if not text.strip():
            raise Exception("No text content found in document")
            
        logger.info(f"✅ Total extracted text: {len(text)} characters")
        return text
        
    except Exception as e:
        error_msg = f"DOCX extraction failed: {str(e)}"
        logger.error(f"❌ {error_msg}")
        raise Exception(error_msg)

def extract_text_from_file(uploaded_file):
    """Extract text based on file type with enhanced error handling"""
    try:
        file_ext = uploaded_file.name.lower().split('.')[-1]
        logger.info(f"\n🔍 Processing file: {uploaded_file.name} (Type: {file_ext.upper()})")
        
        if file_ext == "pdf":
            return extract_text_from_pdf(uploaded_file)
        elif file_ext in ["docx", "doc"]:
            return extract_text_from_docx(uploaded_file)
        else:
            error_msg = f"Unsupported file type: {file_ext}"
            logger.error(f"❌ {error_msg}")
            raise Exception(error_msg)
    except Exception as e:
        logger.error(f"❌ File extraction failed for {uploaded_file.name}: {str(e)}")
        raise

# Improved Claude API functions with seed parameter
def get_resume_extraction_prompt(resume_text, filename_experience=None):
    """Generate enhanced prompt for resume information extraction with consistent experience formatting"""
    experience_instruction = ""
    if filename_experience:
        experience_instruction = f"""
    IMPORTANT - The total work experience for this candidate is "{filename_experience}". 
    Use this exact value for the "Total_Experience" field in your JSON output.
    """
    
    prompt = f"""You are an expert resume information extractor for an Applicant Tracking System. Your task is to extract information exactly as it appears in the resume and format it as clean, structured JSON.

{experience_instruction}

CRITICAL: Your response must be ONLY a valid JSON object with this exact structure (no additional text, explanations, or markdown):

{{
    "Name": "", 
    "Phone": "",  
    "Email": "",    
    "Location": "", 
    "Links": [],    
    "Profile_Summary": "", 
    "Skills": [],   
    "Education": [
        {{
            "Degree": "",      
            "Institution": "", 
            "Year": "",        
            "Score": "",       
            "Type": ""         
        }}
    ],
    "Certifications": [], 
    "Total_Experience": "", 
    "Work_Experience": [
        {{
            "Position": "",     
            "Company": "",      
            "Location": "",     
            "Duration": "",     
            "Start_Date": "",   
            "End_Date": "",     
            "Responsibilities": [] 
        }}
    ],
    "Projects": [
        {{
            "Title": "",        
            "Duration": "",     
            "Description": "",  
            "Technologies": []  
        }}
    ],
    "Additional_Information": "" 
}}

EXTRACTION RULES:
1. Extract information exactly as stated in the resume - do not interpret or modify
2. For dates, preserve the format shown in resume. Use "Present" for current positions
3. For "Phone", extract ONLY phone numbers (remove labels like "Phone:" or "Mobile:")
4. For "Email", extract only email addresses (remove labels)
5. For education, classify Type as "Degree", "12th", "10th", or "Certification"
6. Extract ALL skills mentioned, preserving original terminology
7. For responsibilities, create distinct bullet points as separate array elements
8. Use empty strings ("") for missing text fields and empty arrays ([]) for missing list fields
9. For Skills, include both technical and soft skills as separate items in the array

CRITICAL EXPERIENCE FORMATTING RULES:
10. For "Total_Experience", ALWAYS format as "X years" or "X+ years" (use numbers + "years"):
    - "13+ years" → "13+ years" ✓
    - "9 years" → "9 years" ✓  
    - "Five years" → "5 years" ✓
    - "2-3 years" → "2+ years" ✓
    - "6 months" → "0+ years" ✓
    - "1.5 years" → "1+ years" ✓
    - "Fresher" → "0 years" ✓
    - "Entry level" → "0 years" ✓

11. If experience is unclear from text, calculate from work history dates and format as "X years"
12. Never use decimals, always round down and add "+" if there are additional months
13. Always use the format "NUMBER+ years" or "NUMBER years" - be consistent with this pattern

IMPORTANT: Return ONLY the JSON object. No explanations, no markdown formatting, no additional text.

Resume Text:
{resume_text}"""
    
    return prompt

def get_candidate_scoring_prompt(job_description, candidate_data, weights, additional_preferences=""):
    """Generate enhanced prompt for candidate scoring with better cross-role matching"""
    
    # Format candidate data more robustly
    candidate_skills = candidate_data.get("Skills", [])
    if isinstance(candidate_skills, list):
        candidate_skills_text = ", ".join(candidate_skills) if candidate_skills else "None listed"
    else:
        candidate_skills_text = str(candidate_skills) if candidate_skills else "None listed"
    
    candidate_experience = candidate_data.get("Total_Experience", "Not specified")
    
    # Format work experience with better handling
    work_experience = ""
    if candidate_data.get("Work_Experience") and isinstance(candidate_data["Work_Experience"], list):
        for idx, exp in enumerate(candidate_data["Work_Experience"]):
            if isinstance(exp, dict):
                position = exp.get('Position', 'Unknown')
                company = exp.get('Company', 'Unknown')
                duration = exp.get('Duration', 'Duration not specified')
                work_experience += f"• {position} at {company} ({duration})\n"
                
                responsibilities = exp.get("Responsibilities", [])
                if responsibilities and isinstance(responsibilities, list):
                    work_experience += "  Key Responsibilities:\n"
                    for resp in responsibilities[:3]:  # Limit to top 3 for conciseness
                        work_experience += f"    - {resp}\n"
                work_experience += "\n"
    else:
        work_experience = "No work experience listed"
    
    # Format education
    education = ""
    if candidate_data.get("Education") and isinstance(candidate_data["Education"], list):
        for idx, edu in enumerate(candidate_data["Education"]):
            if isinstance(edu, dict):
                degree = edu.get('Degree', 'Unknown')
                institution = edu.get('Institution', 'Unknown')
                year = edu.get('Year', 'Year not specified')
                education += f"• {degree} from {institution} ({year})\n"
    else:
        education = "No education listed"
    
    # Format certifications
    certifications = ""
    certs = candidate_data.get("Certifications", [])
    if certs:
        if isinstance(certs, list):
            certifications = ", ".join([str(cert) for cert in certs])
        else:
            certifications = str(certs)
    else:
        certifications = "None listed"
    
    additional_criteria = ""
    if additional_preferences.strip():
        additional_criteria = f"""
ADDITIONAL HIRING PREFERENCES:
{additional_preferences}

Consider these preferences when scoring, but prioritize transferable skills and potential over exact matches.
"""
    
    prompt = f"""You are an expert HR consultant specializing in candidate evaluation and cross-functional role assessment. Your task is to evaluate how well a candidate matches a job description, focusing on transferable skills, potential, and adaptability.

SCORING FRAMEWORK:
- Experience Weight: {weights['experience'] * 100}%
- Skills Weight: {weights['skills'] * 100}%
- Education Weight: {weights['education'] * 100}%
- Certifications Weight: {weights['certification'] * 100}%

JOB DESCRIPTION:
{job_description}

{additional_criteria}

CANDIDATE PROFILE:
Name: {candidate_data.get('Name', 'Not provided')}
Total Experience: {candidate_experience}
Skills: {candidate_skills_text}

Work Experience:
{work_experience}

Education:
{education}

Certifications: {certifications}

Profile Summary: {candidate_data.get('Profile_Summary', 'Not provided')}

EVALUATION INSTRUCTIONS:
1. Look beyond exact job title matches - focus on transferable skills and potential
2. Consider how technical skills can translate across domains (e.g., system architecture knowledge valuable for product management)
3. Evaluate leadership, problem-solving, and strategic thinking capabilities
4. Assess learning agility and adaptability for role transitions
5. Score each category 0-100 based on relevance and potential value to the role
6. Be generous with cross-functional skills but honest about gaps

Return ONLY a JSON object with this structure:

{{
    "candidate_match": {{
        "name": "{candidate_data.get('Name', 'Unknown')}",
        "match_details": {{
            "experience": {{
                "score": 85,
                "details": "Detailed explanation focusing on transferable experience and leadership potential"
            }},
            "skills": {{
                "score": 70,
                "matching_skills": ["skill1", "skill2"],
                "missing_skills": ["skill3", "skill4"],
                "transferable_skills": ["skill5", "skill6"],
                "details": "Explanation of how technical/functional skills translate to this role"
            }},
            "education": {{
                "score": 90,
                "details": "Assessment of educational background relevance"
            }},
            "certifications": {{
                "score": 60,
                "details": "Evaluation of certifications and their applicability"
            }}
        }},
        "overall_score": 78.5,
        "explanation": "Comprehensive explanation emphasizing candidate's potential and fit, acknowledging both strengths and development areas",
        "key_strengths": ["strength1", "strength2", "strength3"],
        "key_gaps": ["gap1", "gap2"],
        "recommendation": "HIGHLY_RECOMMENDED"
    }}
}}

RECOMMENDATION GUIDELINES:
- HIGHLY_RECOMMENDED (80-100%): Strong match with excellent transferable skills
- RECOMMENDED (65-79%): Good match with some skill gaps but high potential
- CONSIDER (50-64%): Moderate match, may work with training/development
- NOT_RECOMMENDED (0-49%): Poor match with significant gaps

Focus on potential and transferable value, not just exact matches."""
    
    return prompt

def call_claude_api(client, prompt, max_tokens=3000, seed=None):
    """Call Claude API with enhanced error handling, retry logic, and seed parameter for consistency"""
    if not client:
        logger.error("❌ Claude client not initialized")
        return None
        
    max_retries = 3
    retry_delay = 2
    
    for attempt in range(max_retries):
        try:
            logger.info(f"🤖 Calling Claude API (attempt {attempt + 1}) with seed: {seed}...")
            
            # Prepare the request parameters
            request_params = {
                "model": os.getenv("CLAUDE_MODEL", "claude-3-5-sonnet-20241022"),
                "max_tokens": max_tokens,
                "temperature": 0.1,  # Lower temperature for more consistent JSON output
                "messages": [
                    {"role": "user", "content": prompt}
                ]
            }
            
            # Add seed if provided (for consistency)
            if seed is not None:
                # Note: Anthropic API doesn't have a direct seed parameter like OpenAI
                # We'll use the seed to modify the system message for consistency
                consistent_instruction = f"\nFor consistency, use this processing seed: {seed}"
                request_params["messages"][0]["content"] += consistent_instruction
            
            response = client.messages.create(**request_params)
            logger.info("✅ Claude API call successful")
            return response.content[0].text
            
        except anthropic.RateLimitError as e:
            logger.warning(f"⏳ Rate limit hit (attempt {attempt + 1}): {str(e)}")
            if attempt < max_retries - 1:
                delay = retry_delay * (2 ** attempt)  # Exponential backoff
                logger.info(f"⏳ Retrying in {delay} seconds...")
                time.sleep(delay)
            else:
                logger.error("❌ Rate limit exceeded, all retry attempts failed")
                return None
                
        except anthropic.APIError as e:
            logger.error(f"❌ Claude API error (attempt {attempt + 1}): {str(e)}")
            if attempt < max_retries - 1:
                logger.info(f"⏳ Retrying in {retry_delay} seconds...")
                time.sleep(retry_delay)
                retry_delay *= 2  # Exponential backoff
            else:
                logger.error("❌ All retry attempts failed")
                return None
                
        except Exception as e:
            error_msg = f"Unexpected error (attempt {attempt + 1}): {str(e)}"
            logger.error(f"❌ {error_msg}")
            
            if attempt < max_retries - 1:
                logger.info(f"⏳ Retrying in {retry_delay} seconds...")
                time.sleep(retry_delay)
                retry_delay *= 2  # Exponential backoff
            else:
                logger.error("❌ All retry attempts failed")
                return None
            
def normalize_experience_format(experience_text):
    """Normalize experience format to be consistent across all candidates"""
    if not experience_text or experience_text in ["Unknown", "Not specified", ""]:
        return "Not specified"
    
    experience_text = str(experience_text).strip()
    
    # Convert text numbers to digits
    text_to_num = {
        'zero': '0', 'one': '1', 'two': '2', 'three': '3', 'four': '4', 'five': '5',
        'six': '6', 'seven': '7', 'eight': '8', 'nine': '9', 'ten': '10',
        'eleven': '11', 'twelve': '12', 'thirteen': '13', 'fourteen': '14', 'fifteen': '15'
    }
    
    experience_lower = experience_text.lower()
    
    # Handle fresher/entry level cases
    if any(word in experience_lower for word in ['fresher', 'entry level', 'entry-level', 'graduate', 'no experience']):
        return "0 years"
    
    # Replace text numbers with digits
    for text_num, digit in text_to_num.items():
        experience_lower = experience_lower.replace(text_num, digit)
    
    # Extract numbers from the text
    numbers = re.findall(r'\d+\.?\d*', experience_lower)
    
    if not numbers:
        return experience_text  # Return original if no numbers found
    
    # Get the main number (usually the first or largest)
    main_number = max([float(num) for num in numbers])
    
    # Handle decimal cases
    if main_number < 1:
        if main_number > 0:
            return "0+ years"
        else:
            return "0 years"
    elif main_number == int(main_number):
        # Whole number
        years = int(main_number)
        if '+' in experience_text or 'over' in experience_lower or 'above' in experience_lower or '-' in experience_text:
            return f"{years}+ years"
        else:
            return f"{years} years"
    else:
        # Decimal number - round down and add +
        years = int(main_number)
        return f"{years}+ years"

def parse_json_response(response_text, filename=""):
    """Parse JSON from Claude response with enhanced error handling and experience normalization"""
    try:
        logger.info(f"📝 Parsing JSON response for {filename}...")
        
        # Clean the response text
        response_text = response_text.strip()
        
        # Remove any markdown formatting if present
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        response_text = response_text.strip()
        
        # Try to parse directly first
        result = json.loads(response_text)
        
        # Normalize the experience format for consistency
        if "Total_Experience" in result:
            result["Total_Experience"] = normalize_experience_format(result["Total_Experience"])
        
        logger.info("✅ JSON parsing successful")
        return result
        
    except json.JSONDecodeError as json_error:
        logger.warning(f"⚠️ Direct JSON parsing failed: {str(json_error)}")
        
        # Try to extract JSON from response if it has extra text
        try:
            logger.info("🔍 Attempting to extract JSON from response...")
            # Look for JSON block with more flexible regex
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                result = json.loads(json_str)
                
                # Normalize the experience format for consistency
                if "Total_Experience" in result:
                    result["Total_Experience"] = normalize_experience_format(result["Total_Experience"])
                
                logger.info("✅ JSON extraction successful")
                return result
            else:
                logger.error("❌ No JSON structure found in response")
        except Exception as extraction_error:
            logger.error(f"❌ JSON extraction failed: {str(extraction_error)}")
        
        # Log the raw response for debugging (first 500 chars)
        logger.info(f"🔍 Raw response content:\n{response_text[:500]}...")
        
        # Return enhanced fallback structure
        logger.warning("⚠️ Returning fallback JSON structure")
        return {
            "Name": f"Parse_Failed_{filename}",
            "Phone": "",
            "Email": "",
            "Location": "",
            "Links": [],
            "Profile_Summary": "Failed to parse resume content",
            "Skills": [],
            "Education": [],
            "Certifications": [],
            "Total_Experience": "Not specified",
            "Work_Experience": [],
            "Projects": [],
            "Additional_Information": f"Parsing failed for {filename}"
        }

def extract_experience_from_filename(filename):
    """Extract experience from filename pattern like [2y_6m]"""
    try:
        exp_pattern = r'\[(\d+)y_(\d+)m\]'
        match = re.search(exp_pattern, filename)
        
        if match:
            years = int(match.group(1))
            months = int(match.group(2))
            
            if years == 0 and months == 0:
                return "No Experience"
            elif years == 0:
                return f"{months} months"
            elif months == 0:
                return f"{years} years"
            else:
                return f"{years} years and {months} months"
        
        return None
    except Exception as e:
        logger.warning(f"Could not extract experience from filename {filename}: {e}")
        return None

def process_resume_batch(uploaded_files, client, seed=None):
    """Process uploaded resumes in batches with improved error handling"""
    successful_resumes = []
    failed_count = 0
    
    logger.info(f"\n🚀 Starting batch processing of {len(uploaded_files)} files with seed: {seed}")
    logger.info("=" * 80)
    
    # Create progress bar
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, uploaded_file in enumerate(uploaded_files):
        logger.info(f"\n📁 Processing file {i+1}/{len(uploaded_files)}: {uploaded_file.name}")
        status_text.text(f"Extracting data from {uploaded_file.name}...")
        progress_bar.progress((i + 1) / len(uploaded_files))
        
        try:
            # Extract text from file
            logger.info("📄 Starting text extraction...")
            text = extract_text_from_file(uploaded_file)
            
            if not text.strip():
                logger.error(f"❌ No text content extracted from {uploaded_file.name}")
                failed_count += 1
                continue
            
            logger.info(f"✅ Text extraction successful. Length: {len(text)} characters")
            
            # Extract experience from filename if available
            experience = extract_experience_from_filename(uploaded_file.name)
            if experience:
                logger.info(f"📅 Experience from filename: {experience}")
            
            # Generate prompt
            logger.info("🔧 Generating extraction prompt...")
            prompt = get_resume_extraction_prompt(text, experience)
            
            # Call Claude API with seed for consistency
            logger.info("🤖 Calling Claude API for data extraction...")
            response = call_claude_api(client, prompt, seed=seed)
            
            if not response:
                logger.error(f"❌ Claude API call failed for {uploaded_file.name}")
                failed_count += 1
                continue
            
            logger.info("✅ Claude API response received")
            
            # Parse JSON response
            logger.info("📝 Parsing JSON response...")
            candidate_data = parse_json_response(response, uploaded_file.name)
            
            # Validate and clean the extracted data
            if not candidate_data.get("Name") or candidate_data.get("Name").startswith("Parse_Failed_"):
                logger.warning(f"⚠️ Parsing issues detected for {uploaded_file.name}, but continuing with available data")
            
            # Add metadata
            candidate_data["Source_File"] = uploaded_file.name
            candidate_data["Extraction_Date"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            candidate_data["Processing_Seed"] = seed
            
            # Print extracted data summary
            logger.info(f"\n=== EXTRACTED RESUME DATA FOR {uploaded_file.name} ===")
            logger.info(f"Name: {candidate_data.get('Name', 'N/A')}")
            logger.info(f"Email: {candidate_data.get('Email', 'N/A')}")
            logger.info(f"Phone: {candidate_data.get('Phone', 'N/A')}")
            logger.info(f"Experience: {candidate_data.get('Total_Experience', 'N/A')}")
            logger.info(f"Skills Count: {len(candidate_data.get('Skills', []))}")
            logger.info(f"Education Records: {len(candidate_data.get('Education', []))}")
            logger.info("=" * 60)
            
            successful_resumes.append(candidate_data)
            logger.info(f"✅ Successfully processed {uploaded_file.name}")
            
        except Exception as e:
            logger.error(f"❌ Error processing {uploaded_file.name}: {str(e)}")
            logger.error(f"📍 Traceback: {traceback.format_exc()}")
            failed_count += 1
            
            # Show error in UI but continue processing
            st.error(f"⚠️ Failed to process {uploaded_file.name}: {str(e)}")
    
    # Update session state
    st.session_state.successful_count = len(successful_resumes)
    st.session_state.failed_count = failed_count
    
    logger.info(f"\n🎉 Batch processing completed!")
    logger.info(f"✅ Successful: {len(successful_resumes)}")
    logger.info(f"❌ Failed: {failed_count}")
    
    status_text.text("Data extraction completed!")
    return successful_resumes

def score_candidates_in_batches(candidates, job_description, client, weights, additional_preferences="", batch_size=3, seed=None):
    """Score candidates in smaller batches with improved error handling and seed for consistency"""
    scored_candidates = []
    
    logger.info(f"\n🎯 Starting candidate scoring for {len(candidates)} candidates with seed: {seed}")
    
    # Create smaller batches for better reliability
    batches = [candidates[i:i + batch_size] for i in range(0, len(candidates), batch_size)]
    
    total_batches = len(batches)
    batch_progress = st.progress(0)
    batch_status = st.empty()
    
    for batch_idx, batch in enumerate(batches):
        logger.info(f"\n📊 Processing batch {batch_idx + 1}/{total_batches}")
        batch_status.text(f"Analyzing batch {batch_idx + 1} of {total_batches} against job requirements...")
        batch_progress.progress((batch_idx + 1) / total_batches)
        
        for candidate in batch:
            candidate_name = candidate.get('Name', 'Unknown')
            logger.info(f"🔍 Scoring candidate: {candidate_name}")
            
            try:
                prompt = get_candidate_scoring_prompt(job_description, candidate, weights, additional_preferences)
                response = call_claude_api(client, prompt, max_tokens=2500, seed=seed)
                
                if response:
                    match_data = parse_json_response(response, candidate_name)
                    
                    if "candidate_match" in match_data:
                        candidate_match = match_data["candidate_match"]
                        
                        candidate_record = {
                            'candidate_data': candidate,
                            'match_details': candidate_match.get('match_details', {}),
                            'overall_score': candidate_match.get('overall_score', 0),
                            'explanation': candidate_match.get('explanation', "No explanation provided"),
                            'key_strengths': candidate_match.get('key_strengths', []),
                            'key_gaps': candidate_match.get('key_gaps', []),
                            'recommendation': candidate_match.get('recommendation', 'CONSIDER'),
                            'processing_seed': seed
                        }
                        
                        scored_candidates.append(candidate_record)
                        logger.info(f"✅ Scored {candidate_name}: {candidate_record['overall_score']:.1f}%")
                    else:
                        logger.warning(f"⚠️ Invalid response structure for {candidate_name}, using fallback")
                        scored_candidates.append(create_fallback_score(candidate, seed))
                else:
                    logger.error(f"❌ No API response for {candidate_name}, using fallback")
                    scored_candidates.append(create_fallback_score(candidate, seed))
                    
            except Exception as e:
                logger.error(f"❌ Error scoring {candidate_name}: {str(e)}")
                scored_candidates.append(create_fallback_score(candidate, seed))
                # Show error in UI but continue processing
                st.warning(f"⚠️ Failed to score {candidate_name}: {str(e)}")
        
        # Longer delay between batches to avoid rate limiting
        if batch_idx < total_batches - 1:  # Don't delay after the last batch
            time.sleep(3)
    
    batch_status.text("Candidate scoring completed!")
    
    # Sort by score
    scored_candidates.sort(key=lambda x: x['overall_score'], reverse=True)
    if scored_candidates:
        logger.info(f"🏆 Scoring completed. Top candidate: {scored_candidates[0]['candidate_data'].get('Name', 'Unknown')} ({scored_candidates[0]['overall_score']:.1f}%)")
    
    return scored_candidates

def create_fallback_score(candidate, seed=None):
    """Create fallback score when API fails"""
    return {
        'candidate_data': candidate,
        'match_details': {
            'experience': {'score': 50, 'details': 'Could not evaluate due to processing issues'},
            'skills': {'score': 50, 'details': 'Could not evaluate due to processing issues', 'matching_skills': [], 'missing_skills': []},
            'education': {'score': 50, 'details': 'Could not evaluate due to processing issues'},
            'certifications': {'score': 50, 'details': 'Could not evaluate due to processing issues'}
        },
        'overall_score': 50,
        'explanation': "Could not fully evaluate due to processing issues. Manual review recommended.",
        'key_strengths': ["Unable to determine - manual review needed"],
        'key_gaps': ["Unable to determine - manual review needed"],
        'recommendation': 'CONSIDER',
        'processing_seed': seed
    }

def safe_convert_to_string(value):
    """Safely convert any value to string, handling lists and dicts"""
    try:
        if value is None:
            return ""
        elif isinstance(value, str):
            return value
        elif isinstance(value, list):
            # Handle list of strings
            if all(isinstance(item, str) for item in value):
                return ", ".join(value)
            # Handle list of dicts or mixed types
            else:
                result = []
                for item in value:
                    if isinstance(item, dict):
                        # Convert dict to readable string
                        dict_str = "; ".join([f"{k}: {v}" for k, v in item.items() if v])
                        result.append(dict_str)
                    else:
                        result.append(str(item))
                return " | ".join(result)
        elif isinstance(value, dict):
            # Convert dict to readable string
            return "; ".join([f"{k}: {v}" for k, v in value.items() if v])
        else:
            return str(value)
    except Exception as e:
        logger.warning(f"Error converting value to string: {e}")
        return str(value) if value else ""

def convert_to_dataframe(resumes_data):
    """Convert resume data to DataFrame with enhanced error handling"""
    if not resumes_data:
        return None
    
    try:
        # Create DataFrame
        df = pd.DataFrame(resumes_data)
        
        # Process nested fields safely
        def extract_education(edu_list):
            try:
                if not edu_list or not isinstance(edu_list, list):
                    return ""
                result = []
                for e in edu_list:
                    if isinstance(e, dict):
                        degree = e.get('Degree', '').strip()
                        institution = e.get('Institution', '').strip()
                        year = e.get('Year', '').strip()
                        if degree or institution:
                            edu_str = f"{degree} from {institution} ({year})".strip()
                            result.append(edu_str)
                return "; ".join([r for r in result if r and r != " from  ()"])
            except Exception as e:
                logger.warning(f"Error extracting education: {e}")
                return ""
        
        def extract_work_exp(exp_list):
            try:
                if not exp_list or not isinstance(exp_list, list):
                    return ""
                result = []
                for e in exp_list:
                    if isinstance(e, dict):
                        position = e.get('Position', '').strip()
                        company = e.get('Company', '').strip()
                        duration = e.get('Duration', '').strip()
                        if position or company:
                            exp_str = f"{position} at {company} ({duration})".strip()
                            result.append(exp_str)
                return "; ".join([r for r in result if r and r != " at  ()"])
            except Exception as e:
                logger.warning(f"Error extracting work experience: {e}")
                return ""
        
        def extract_skills(skills_list):
            try:
                if not skills_list or not isinstance(skills_list, list):
                    return ""
                return ", ".join([str(skill).strip() for skill in skills_list if str(skill).strip()])
            except Exception as e:
                logger.warning(f"Error extracting skills: {e}")
                return ""
        
        def extract_projects(proj_list):
            try:
                if not proj_list or not isinstance(proj_list, list):
                    return ""
                result = []
                for p in proj_list:
                    if isinstance(p, dict):
                        title = p.get('Title', '').strip()
                        description = p.get('Description', '').strip()
                        if title:
                            proj_str = f"{title}: {description}".strip()
                            result.append(proj_str)
                return "; ".join([r for r in result if r and not r.endswith(': ')])
            except Exception as e:
                logger.warning(f"Error extracting projects: {e}")
                return ""
        
        # Apply processing to nested fields
        try:
            if "Education" in df.columns:
                df["Education_Summary"] = df["Education"].apply(extract_education)
                df = df.drop(columns=["Education"])
                
            if "Work_Experience" in df.columns:
                df["Work_Summary"] = df["Work_Experience"].apply(extract_work_exp)
                df = df.drop(columns=["Work_Experience"])
                
            if "Skills" in df.columns:
                df["Skills_List"] = df["Skills"].apply(extract_skills)
                df = df.drop(columns=["Skills"])
            
            if "Projects" in df.columns:
                df["Projects_Summary"] = df["Projects"].apply(extract_projects)
                df = df.drop(columns=["Projects"])
        
        except Exception as e:
            logger.error(f"Error processing nested fields: {e}")
        
        # Convert any remaining complex columns to strings safely
        for col in df.columns:
            if df[col].dtype == 'object':
                df[col] = df[col].apply(safe_convert_to_string)
        
        return df
        
    except Exception as e:
        logger.error(f"Error converting to dataframe: {e}")
        st.error(f"Error processing candidate data: {e}")
        return None

def create_excel_report(scored_candidates, job_description):
    """Create comprehensive Excel report with enhanced formatting"""
    try:
        output = BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Main candidates sheet
            main_data = []
            for idx, candidate in enumerate(scored_candidates):
                try:
                    candidate_info = candidate['candidate_data']
                    match_details = candidate['match_details']
                    
                    # Safely extract skills
                    skills = candidate_info.get("Skills", [])
                    if isinstance(skills, list):
                        skills_text = ", ".join(skills)
                    else:
                        skills_text = str(skills)
                    
                    # Safely extract matching and missing skills
                    matching_skills = match_details.get('skills', {}).get('matching_skills', [])
                    missing_skills = match_details.get('skills', {}).get('missing_skills', [])
                    
                    row = {
                        "Rank": idx + 1,
                        "Name": candidate_info.get("Name", "Unknown"),
                        "Overall_Score": f"{candidate.get('overall_score', 0):.1f}%",
                        "Recommendation": candidate.get('recommendation', 'CONSIDER'),
                        "Experience_Score": f"{match_details.get('experience', {}).get('score', 0):.1f}%",
                        "Skills_Score": f"{match_details.get('skills', {}).get('score', 0):.1f}%",
                        "Education_Score": f"{match_details.get('education', {}).get('score', 0):.1f}%",
                        "Certification_Score": f"{match_details.get('certifications', {}).get('score', 0):.1f}%",
                        "Total_Experience": candidate_info.get("Total_Experience", "Not specified"),
                        "Phone": candidate_info.get("Phone", ""),
                        "Email": candidate_info.get("Email", ""),
                        "Location": candidate_info.get("Location", ""),
                        "Skills": skills_text[:500] if len(skills_text) > 500 else skills_text,  # Limit length for Excel
                        "Matching_Skills": ", ".join(matching_skills) if matching_skills else "",
                        "Missing_Skills": ", ".join(missing_skills) if missing_skills else "",
                        "Key_Strengths": ", ".join(candidate.get('key_strengths', [])),
                        "Key_Gaps": ", ".join(candidate.get('key_gaps', [])),
                        "Explanation": candidate.get('explanation', '')[:500],  # Limit for Excel
                        "Resume_File": candidate_info.get("Source_File", ""),
                        "Processing_Seed": candidate.get('processing_seed', 'N/A')
                    }
                    main_data.append(row)
                except Exception as e:
                    logger.warning(f"Error processing candidate {idx}: {e}")
                    continue
            
            main_df = pd.DataFrame(main_data)
            main_df.to_excel(writer, sheet_name='All_Candidates', index=False)
            
            # Top 10 candidates sheet
            top_10 = main_df.head(10)
            top_10.to_excel(writer, sheet_name='Top_10_Candidates', index=False)
            
            # Job description sheet
            job_data = [
                {"Field": "Job Description", "Content": job_description[:1000]},  # Limit length
                {"Field": "Generated On", "Content": datetime.now().strftime("%Y-%m-%d %H:%M:%S")},
                {"Field": "Total Candidates", "Content": len(scored_candidates)},
                {"Field": "Highly Recommended", "Content": len([c for c in scored_candidates if c.get('recommendation') == 'HIGHLY_RECOMMENDED'])},
                {"Field": "Recommended", "Content": len([c for c in scored_candidates if c.get('recommendation') == 'RECOMMENDED'])},
            ]
            job_df = pd.DataFrame(job_data)
            job_df.to_excel(writer, sheet_name='Job_Info', index=False)
        
        return output.getvalue()
        
    except Exception as e:
        logger.error(f"Error creating Excel report: {e}")
        st.error(f"Error creating Excel report: {e}")
        return None
    
def check_password():
    """Returns `True` if the user had the correct password."""
    
    def password_entered():
        """Checks whether a password entered by the user is correct."""
        try:
            username = st.session_state["username"].strip().lower()
            password = st.session_state["password"]
            
            # Check if username exists and password matches
            if username in VALID_USERS and VALID_USERS[username] == password:
                st.session_state["password_correct"] = True
                st.session_state["authenticated_user"] = username
                del st.session_state["password"]  # Don't store password
                del st.session_state["username"]  # Don't store username
            else:
                st.session_state["password_correct"] = False
                st.session_state["login_attempted"] = True  # Track that login was attempted
        except Exception as e:
            logger.error(f"Error in password validation: {e}")
            st.session_state["password_correct"] = False
            st.session_state["login_attempted"] = True

    # Return True if password is validated
    if st.session_state.get("password_correct", False):
        return True

    # Show login form
    st.markdown("""
    <div style="max-width: 400px; margin: 50px auto; padding: 2rem; 
                background: white; border-radius: 10px; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);">
        <h2 style="text-align: center; color: #333; margin-bottom: 2rem;">
            🔐 SmartWorks Login
        </h2>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.text_input(
            "👤 Username", 
            key="username", 
            placeholder="Enter your username",
            help="Use: smartworks_admin, client_manager, operations, or ansh.arora1@sworks.co.in"
        )
        st.text_input(
            "🔑 Password", 
            type="password", 
            key="password",
            placeholder="Enter your password"
        )
        
        if st.button("🚀 Login", use_container_width=True, type="primary"):
            password_entered()
        
        # Show error message ONLY if login was attempted and failed
        if (st.session_state.get("login_attempted", False) and 
            not st.session_state.get("password_correct", False)):
            st.error("❌ Invalid username or password")
        
        # Add some styling
        st.markdown("""
        <div style="text-align: center; margin-top: 2rem; color: #666;">
            <small>Access restricted to authorized SmartWorks personnel only</small>
        </div>
        """, unsafe_allow_html=True)

    return False

def get_user_role(username):
    """Get user role based on username"""
    role_mapping = {
        "smartworks_admin": "Admin",
        "client_manager": "Client Manager", 
        "operations": "Operations",
        "ansh.arora1@sworks.co.in": "Developer"
    }
    return role_mapping.get(username, "User")

def show_user_info():
    """Display logged in user info in sidebar"""
    try:
        if "authenticated_user" in st.session_state:
            username = st.session_state["authenticated_user"]
            role = get_user_role(username)
            
            st.sidebar.markdown("---")
            st.sidebar.markdown("### 👤 User Info")
            st.sidebar.write(f"**User:** {username}")
            st.sidebar.write(f"**Role:** {role}")
            
            if st.sidebar.button("🚪 Logout"):
                # Clear all session state
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.rerun()
    except Exception as e:
        logger.error(f"Error showing user info: {e}")

def configure_scoring_weights():
    """Enhanced weight configuration interface"""
    try:
        st.markdown("### ⚙️ Scoring Configuration")
        
        # Show current weights
        st.markdown("**Current Active Weights:**")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("🎯 Experience", f"{st.session_state.weights['experience']:.0%}")
        with col2:
            st.metric("🛠️ Skills", f"{st.session_state.weights['skills']:.0%}")
        with col3:
            st.metric("🎓 Education", f"{st.session_state.weights['education']:.0%}")
        with col4:
            st.metric("📜 Certifications", f"{st.session_state.weights['certification']:.0%}")
        
        # Quick preset buttons
        st.markdown("---")
        st.markdown("**⚡ Quick Presets:**")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("💼 Experience Focus", use_container_width=True, help="50% Experience, 30% Skills, 15% Education, 5% Certifications"):
                st.session_state.weights = {"experience": 0.50, "skills": 0.30, "education": 0.15, "certification": 0.05}
        
        with col2:
            if st.button("🛠️ Skills Focus", use_container_width=True, help="20% Experience, 60% Skills, 15% Education, 5% Certifications"):
                st.session_state.weights = {"experience": 0.20, "skills": 0.60, "education": 0.15, "certification": 0.05}
        
        with col3:
            if st.button("⚖️ Balanced", use_container_width=True, help="30% Experience, 40% Skills, 20% Education, 10% Certifications"):
                st.session_state.weights = {"experience": 0.30, "skills": 0.40, "education": 0.20, "certification": 0.10}
        
        st.markdown("---")
        
        # Manual weight adjustment
        st.markdown("**🎚️ Set Custom Weights:**")
        st.markdown("*Set each weight percentage - they will be normalized to 100% when you apply*")
        
        # Create input fields for weights
        col1, col2 = st.columns(2)
        
        with col1:
            exp_weight = st.number_input(
                "🎯 Experience (%)",
                min_value=0,
                max_value=100,
                value=int(st.session_state.weights['experience'] * 100),
                step=5,
                help="Set weight for work experience"
            )
            
            skills_weight = st.number_input(
                "🛠️ Skills (%)",
                min_value=0,
                max_value=100,
                value=int(st.session_state.weights['skills'] * 100),
                step=5,
                help="Set weight for technical and soft skills"
            )
        
        with col2:
            edu_weight = st.number_input(
                "🎓 Education (%)",
                min_value=0,
                max_value=100,
                value=int(st.session_state.weights['education'] * 100),
                step=5,
                help="Set weight for educational background"
            )
            
            cert_weight = st.number_input(
                "📜 Certifications (%)",
                min_value=0,
                max_value=100,
                value=int(st.session_state.weights['certification'] * 100),
                step=5,
                help="Set weight for professional certifications"
            )
        
        # Show total and preview
        total = exp_weight + skills_weight + edu_weight + cert_weight
        
        # Preview what the normalized weights will be
        if total > 0:
            norm_exp = (exp_weight / total) * 100
            norm_skills = (skills_weight / total) * 100
            norm_edu = (edu_weight / total) * 100
            norm_cert = (cert_weight / total) * 100
            
            st.markdown("**📊 Preview (Normalized to 100%):**")
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.info(f"🎯 {norm_exp:.0f}%")
            with col2:
                st.info(f"🛠️ {norm_skills:.0f}%")
            with col3:
                st.info(f"🎓 {norm_edu:.0f}%")
            with col4:
                st.info(f"📜 {norm_cert:.0f}%")
            
            # Apply button
            if st.button("✅ Apply New Weights", type="primary", use_container_width=True):
                # Normalize and apply weights
                st.session_state.weights = {
                    "experience": exp_weight / total,
                    "skills": skills_weight / total,
                    "education": edu_weight / total,
                    "certification": cert_weight / total
                }
                st.success(f"✅ Weights updated! Experience: {norm_exp:.0f}%, Skills: {norm_skills:.0f}%, Education: {norm_edu:.0f}%, Certifications: {norm_cert:.0f}%")
        else:
            st.warning("⚠️ Please set at least one weight above 0")
        
        st.markdown("---")
        st.info("💡 **How it works:** Set your desired percentages above, then click 'Apply New Weights'. The system will automatically balance them to total 100%.")
        
    except Exception as e:
        logger.error(f"Error in configure_scoring_weights: {e}")
        st.error(f"Error configuring weights: {e}")

def display_consistency_info():
    """Display information about processing consistency"""
    try:
        if st.session_state.get('processing_seed'):
            st.sidebar.markdown("---")
            st.sidebar.markdown("### 🎯 Processing Info")
            st.sidebar.write(f"**Seed:** {st.session_state.processing_seed}")
            st.sidebar.info("💡 Same seed = consistent results for same inputs")
    except Exception as e:
        logger.error(f"Error displaying consistency info: {e}")

# Streamlit UI
def main():
    try:
        st.set_page_config(
            page_title="SmartWorks Resume Shortlisting Tool",
            page_icon="🎯",
            layout="wide"
        )
        
        # Initialize session state
        init_session_state()
        
        # Check authentication FIRST
        if not check_password():
            st.stop()
        
        # Enhanced CSS for professional appearance
        st.markdown("""
        <style>
        .main-header {
            text-align: center;
            padding: 1.5rem 0;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border-radius: 15px;
            margin-bottom: 2rem;
            box-shadow: 0 4px 15px rgba(102, 126, 234, 0.3);
        }
        
        .upload-section {
            background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
            padding: 1rem;
            border-radius: 15px;
            margin-bottom: 1rem;
            border: 1px solid #dee2e6;
            box-shadow: 0 2px 10px rgba(0,0,0,0.05);
        }
        
        .top-candidate-card {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border-radius: 20px;
            padding: 2.5rem;
            margin: 1.5rem 0;
            text-align: center;
            box-shadow: 0 8px 30px rgba(102, 126, 234, 0.4);
            position: relative;
            overflow: hidden;
        }
        
        .rank-badge {
            background: rgba(255,255,255,0.2);
            border: 3px solid white;
            border-radius: 50%;
            width: 70px;
            height: 70px;
            display: flex;
            align-items: center;
            justify-content: center;
            margin: 0 auto 1.5rem auto;
            font-size: 1.8rem;
            font-weight: bold;
            backdrop-filter: blur(10px);
            position: relative;
            z-index: 1;
        }
        
        .metric-card {
            background: white;
            padding: 1.5rem;
            border-radius: 12px;
            text-align: center;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            margin: 0.5rem;
            transition: transform 0.3s ease;
        }
        
        .score-badge {
            display: inline-block;
            padding: 0.6rem 1.2rem;
            border-radius: 25px;
            font-weight: bold;
            margin: 0.5rem;
            font-size: 0.9rem;
            position: relative;
            z-index: 1;
        }
        
        .score-excellent { 
            background: linear-gradient(135deg, #10b981 0%, #059669 100%);
            color: white;
            box-shadow: 0 4px 15px rgba(16, 185, 129, 0.3);
        }
        .score-good { 
            background: linear-gradient(135deg, #3b82f6 0%, #2563eb 100%);
            color: white;
            box-shadow: 0 4px 15px rgba(59, 130, 246, 0.3);
        }
        .score-average { 
            background: linear-gradient(135deg, #f59e0b 0%, #d97706 100%);
            color: white;
            box-shadow: 0 4px 15px rgba(245, 158, 11, 0.3);
        }
        .score-poor { 
            background: linear-gradient(135deg, #ef4444 0%, #dc2626 100%);
            color: white;
            box-shadow: 0 4px 15px rgba(239, 68, 68, 0.3);
        }
        </style>
        """, unsafe_allow_html=True)
        
        # Main header with enhanced styling
        st.markdown("""
        <div class="main-header">
            <h1>🎯 SmartWorks AI Resume Shortlisting Tool</h1>
            <p style="font-size: 1.1rem; margin-bottom: 0;">Intelligent candidate screening powered by Advanced AI</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Initialize Claude client
        client = init_claude_client()
        
        if not client:
            st.error("❌ Failed to initialize AI client. Please check your API configuration.")
            st.stop()
        
        # Sidebar for configuration
        with st.sidebar:
            # Show user info
            show_user_info()
            
            # Configure scoring weights
            configure_scoring_weights()
            
            # Display consistency info
            display_consistency_info()

        # Create tabs
        tab1, tab2, tab3 = st.tabs(["📁 Upload & Process", "👥 Candidate Details", "🏆 Shortlisted Candidates"])
        
        with tab1:
            # Upload section with enhanced styling
            st.markdown('<div class="upload-section">', unsafe_allow_html=True)
            st.markdown("### 📄 Upload Resume Files")
            st.markdown("*Upload multiple PDF, DOC, or DOCX resume files for AI-powered analysis*")
            
            uploaded_files = st.file_uploader(
                "Choose Resume Files",
                type=["pdf", "docx", "doc"],
                accept_multiple_files=True,
                help="Upload multiple resume files to process. Supported formats: PDF, DOC, DOCX"
            )
            
            if uploaded_files:
                st.success(f"✅ {len(uploaded_files)} files uploaded successfully")
                
                # Show file details with enhanced styling
                file_details = []
                total_size = 0
                for file in uploaded_files:
                    size_kb = file.size / 1024
                    total_size += size_kb
                    file_details.append({
                        "📄 Filename": file.name,
                        "📋 Type": file.name.split('.')[-1].upper(),
                        "📊 Size": f"{size_kb:.1f} KB"
                    })
                
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("Total Files", len(uploaded_files))
                with col2:
                    st.metric("Total Size", f"{total_size:.1f} KB")
                
                st.dataframe(pd.DataFrame(file_details), use_container_width=True)
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Job requirements section
            st.markdown("### 📋 Job Requirements")
            st.markdown("*Provide detailed job information for accurate candidate matching*")
            
            col1, col2 = st.columns([1, 2])
            with col1:
                job_title = st.text_input(
                    "Job Title",
                    placeholder="e.g., Senior Software Engineer",
                    help="Enter the specific job title"
                )
            
            job_description = st.text_area(
                "Job Description",
                height=200,
                placeholder="Paste the complete job description including requirements, responsibilities, and qualifications...",
                help="Provide comprehensive job description for better AI matching accuracy"
            )
            
            additional_preferences = st.text_area(
                "Additional Hiring Preferences (Optional)",
                height=100,
                placeholder="e.g., Prefer candidates with startup experience, remote work capability, specific certifications...",
                help="Add specific preferences beyond the standard job description"
            )
            
            # Consistency check section
            st.markdown("---")
            st.markdown("### 🔄 Consistency Testing")
            
            # Check if we can do consistency testing
            can_test_consistency = (
                st.session_state.get('current_job_description') == job_description and 
                st.session_state.get('processing_seed') and
                st.session_state.get('successful_resumes')
            )
            
            if can_test_consistency:
                st.info(f"🎯 **Consistency Mode Available**: Same job description detected. Using seed: {st.session_state.processing_seed}")
                st.markdown("*This will process with the same parameters to test result consistency*")
            else:
                st.info("🆕 **New Processing**: This will generate a new processing seed for consistent results")
            
            # Process button with enhanced styling
            st.markdown("---")
            if st.button("🚀 Start AI Processing", type="primary", use_container_width=True):
                if not uploaded_files:
                    st.error("❌ Please upload at least one resume file")
                elif not job_description:
                    st.error("❌ Please enter a job description")
                else:
                    try:
                        # Generate or reuse seed for consistency
                        if can_test_consistency:
                            processing_seed = st.session_state.processing_seed
                            st.info(f"🔄 **Consistency Test**: Using existing seed {processing_seed}")
                        else:
                            processing_seed = generate_processing_seed(job_description, st.session_state.weights)
                            st.session_state.processing_seed = processing_seed
                            st.session_state.current_job_description = job_description
                            st.info(f"🆕 **New Processing**: Generated seed {processing_seed}")
                        
                        # Clear previous results
                        st.session_state.processing_complete = False
                        
                        st.markdown("### ⚡ AI Processing Results")
                        
                        # Step 1: Extract candidate information
                        st.markdown("**📊 AI Resume Data Extraction**")
                        successful_resumes = process_resume_batch(uploaded_files, client, seed=processing_seed)
                        
                        st.session_state.successful_resumes = successful_resumes
                        st.session_state.successful_count = len(successful_resumes)
                        st.session_state.failed_count = len(uploaded_files) - len(successful_resumes)
                        
                        # Show extraction results with metrics
                        col1, col2, col3, col4 = st.columns(4)
                        with col1:
                            st.metric("📁 Total Files", len(uploaded_files))
                        with col2:
                            st.metric("✅ Processed", len(successful_resumes))
                        with col3:
                            st.metric("❌ Failed", st.session_state.failed_count)
                        with col4:
                            success_rate = (len(successful_resumes)/len(uploaded_files)*100) if uploaded_files else 0
                            st.metric("📈 Success Rate", f"{success_rate:.1f}%")
                        
                        if successful_resumes:
                            # Step 2: Score candidates
                            st.markdown("**🎯 AI Candidate Scoring & Ranking**")
                            scored_candidates = score_candidates_in_batches(
                                successful_resumes, 
                                job_description, 
                                client, 
                                st.session_state.weights,
                                additional_preferences,
                                seed=processing_seed
                            )
                            
                            st.session_state.top_candidates = scored_candidates
                            st.session_state.current_job_title = job_title or "Position"
                            st.session_state.processing_complete = True
                            
                            # Show scoring results with enhanced metrics
                            if scored_candidates:
                                highly_recommended = len([c for c in scored_candidates if c.get('recommendation') == 'HIGHLY_RECOMMENDED'])
                                recommended = len([c for c in scored_candidates if c.get('recommendation') == 'RECOMMENDED'])
                                highest_score = max([c['overall_score'] for c in scored_candidates])
                                
                                st.markdown("### 🎉 Processing Complete!")
                                col1, col2, col3, col4 = st.columns(4)
                                with col1:
                                    st.metric("🌟 Highly Recommended", highly_recommended)
                                with col2:
                                    st.metric("👍 Recommended", recommended)
                                with col3:
                                    st.metric("🎯 Processing Seed", processing_seed)
                                with col4:
                                    st.metric("🏆 Top Score", f"{highest_score:.1f}%")
                                
                                st.success("✅ AI processing completed successfully!")
                                
                                # Show consistency info if applicable
                                if can_test_consistency:
                                    st.info("🔄 **Consistency Test Complete**: Compare results with previous runs using the same seed to verify consistency")
                                
                                # Navigation guidance
                                st.markdown("---")
                                st.markdown("### 📋 What's Next?")
                                col1, col2 = st.columns(2)
                                
                                with col1:
                                    st.info("""
                                    **👥 View All Candidates**
                                    
                                    Go to **'Candidate Details'** tab to:
                                    - See all extracted resume data
                                    - Download candidate information
                                    - Review processing results
                                    """)
                                
                                with col2:
                                    st.info("""
                                    **🏆 View Rankings**
                                    
                                    Go to **'Shortlisted Candidates'** tab to:
                                    - See AI-ranked top candidates
                                    - View detailed scoring analysis
                                    - Download comprehensive reports
                                    """)
                            else:
                                st.error("❌ No candidates could be scored. Please check your files and try again.")
                        else:
                            st.error("❌ No resumes were successfully processed. Please check your files and try again.")
                            
                    except Exception as e:
                        logger.error(f"Error in processing: {e}")
                        st.error(f"❌ Processing failed: {e}")
                        st.info("💡 The application will continue running. Please try again or contact support.")
        
        with tab2:
            try:
                st.markdown("### 👥 All Extracted Candidates")
                
                if st.session_state.successful_resumes:
                    # Create and display candidate DataFrame
                    candidate_df = convert_to_dataframe(st.session_state.successful_resumes)
                    st.session_state.candidate_df = candidate_df
                    
                    if candidate_df is not None:
                        # Show summary metrics
                        col1, col2, col3, col4 = st.columns(4)
                        col1.metric("Total Candidates", len(candidate_df))
                        
                        # Calculate unique skills more safely
                        try:
                            all_skills = []
                            for skills in candidate_df.get('Skills_List', []):
                                if isinstance(skills, str) and skills:
                                    all_skills.extend([skill.strip() for skill in skills.split(',') if skill.strip()])
                            unique_skills = len(set(all_skills)) if all_skills else 0
                            col2.metric("Unique Skills", unique_skills)
                        except Exception as e:
                            logger.warning(f"Error calculating skills metrics: {e}")
                            col2.metric("Unique Skills", "N/A")
                        
                        # Experience range calculation
                        try:
                            exp_values = candidate_df.get('Total_Experience', [])
                            exp_count = sum(1 for exp in exp_values if exp and str(exp) != 'Unknown' and str(exp) != 'Not specified')
                            col3.metric("With Experience Info", exp_count)
                        except Exception as e:
                            logger.warning(f"Error calculating experience metrics: {e}")
                            col3.metric("With Experience Info", "N/A")
                        
                        # Education levels
                        try:
                            edu_values = candidate_df.get('Education_Summary', [])
                            edu_count = sum(1 for edu in edu_values if edu and str(edu).strip())
                            col4.metric("With Education Info", edu_count)
                        except Exception as e:
                            logger.warning(f"Error calculating education metrics: {e}")
                            col4.metric("With Education Info", "N/A")
                        
                        st.dataframe(candidate_df, use_container_width=True, height=400)
                        
                        # Download buttons with enhanced styling
                        st.markdown("### 📥 Download Options")
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            try:
                                # Download as Excel
                                excel_buffer = BytesIO()
                                with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                                    candidate_df.to_excel(writer, sheet_name='Candidates', index=False)
                                
                                st.download_button(
                                    "📊 Download Excel Report",
                                    data=excel_buffer.getvalue(),
                                    file_name=f"candidates_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    use_container_width=True,
                                    key="candidates_excel_download"
                                )
                            except Exception as e:
                                logger.error(f"Error creating Excel download: {e}")
                                st.error("Error creating Excel file")
                        
                        with col2:
                            try:
                                # Download as CSV
                                csv_data = candidate_df.to_csv(index=False).encode('utf-8')
                                st.download_button(
                                    "📄 Download CSV Data",
                                    data=csv_data,
                                    file_name=f"candidates_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                    mime="text/csv",
                                    use_container_width=True,
                                    key="candidates_csv_download"
                                )
                            except Exception as e:
                                logger.error(f"Error creating CSV download: {e}")
                                st.error("Error creating CSV file")
                        
                        with col3:
                            try:
                                # Download as JSON
                                json_data = json.dumps(st.session_state.successful_resumes, indent=2)
                                st.download_button(
                                    "🔧 Download JSON Data",
                                    data=json_data,
                                    file_name=f"candidates_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                                    mime="application/json",
                                    use_container_width=True,
                                    key="candidates_json_download"
                                )
                            except Exception as e:
                                logger.error(f"Error creating JSON download: {e}")
                                st.error("Error creating JSON file")
                    else:
                        st.error("❌ Error processing candidate data")
                else:
                    st.info("📝 No candidate data available. Please process some resumes first in the 'Upload & Process' tab.")
            
            except Exception as e:
                logger.error(f"Error in candidate details tab: {e}")
                st.error(f"❌ Error displaying candidate details: {e}")
        
        with tab3:
            try:
                st.markdown("### 🏆 AI-Ranked Top Candidates")
                
                if st.session_state.top_candidates:
                    # Overall statistics
                    total_candidates = len(st.session_state.top_candidates)
                    highly_recommended = len([c for c in st.session_state.top_candidates if c.get('recommendation') == 'HIGHLY_RECOMMENDED'])
                    recommended = len([c for c in st.session_state.top_candidates if c.get('recommendation') == 'RECOMMENDED'])
                    highest_score = max([c['overall_score'] for c in st.session_state.top_candidates])
                    
                    # Show processing seed info
                    processing_seed = st.session_state.get('processing_seed', 'N/A')
                    
                    # Statistics cards
                    st.markdown("**📊 Screening Summary**")
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.markdown(f"""
                        <div class="metric-card">
                            <h3 style="color: #667eea; margin: 0;">👥 {total_candidates}</h3>
                            <p style="margin: 0; color: #6c757d;">Total Analyzed</p>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with col2:
                        st.markdown(f"""
                        <div class="metric-card">
                            <h3 style="color: #10b981; margin: 0;">🌟 {highly_recommended}</h3>
                            <p style="margin: 0; color: #6c757d;">Highly Recommended</p>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with col3:
                        st.markdown(f"""
                        <div class="metric-card">
                            <h3 style="color: #3b82f6; margin: 0;">👍 {recommended}</h3>
                            <p style="margin: 0; color: #6c757d;">Recommended</p>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with col4:
                        st.markdown(f"""
                        <div class="metric-card">
                            <h3 style="color: #f59e0b; margin: 0;">🏆 {highest_score:.1f}%</h3>
                            <p style="margin: 0; color: #6c757d;">Top Score</p>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    # Show seed info
                    st.info(f"🎯 **Processing Seed:** {processing_seed} | *Same seed ensures consistent results for identical inputs*")
                    
                    st.markdown("---")
                    
                    # Top 5 candidates with enhanced cards
                    st.markdown("### 🥇 Top 5 Candidates")
                    top_5 = st.session_state.top_candidates[:5]
                    
                    for idx, candidate in enumerate(top_5):
                        try:
                            candidate_info = candidate['candidate_data']
                            score = candidate['overall_score']
                            match_details = candidate['match_details']
                            
                            # Determine score class and emoji
                            if score >= 85:
                                score_class = "score-excellent"
                                score_emoji = "🎯"
                            elif score >= 70:
                                score_class = "score-good"
                                score_emoji = "✅"
                            elif score >= 55:
                                score_class = "score-average"
                                score_emoji = "⚠️"
                            else:
                                score_class = "score-poor"
                                score_emoji = "❌"
                            
                            # Recommendation styling
                            rec_emoji = {
                                'HIGHLY_RECOMMENDED': '🌟',
                                'RECOMMENDED': '👍',
                                'CONSIDER': '🤔',
                                'NOT_RECOMMENDED': '❌'
                            }.get(candidate.get('recommendation', 'CONSIDER'), '🤔')
                            
                            st.markdown(f"""
                            <div class="top-candidate-card">
                                <div class="rank-badge">#{idx+1}</div>
                                <h2 style="margin-bottom: 1rem; position: relative; z-index: 1;">{candidate_info.get('Name', 'Unknown')}</h2>
                                <div class="score-badge {score_class}">
                                    {score_emoji} {score:.1f}% Match Score
                                </div>
                                <p style="font-size: 1.1rem; margin: 1rem 0;"><strong>📧</strong> {candidate_info.get('Email', 'N/A')}</p>
                                <p style="font-size: 1.1rem; margin: 1rem 0;"><strong>📱</strong> {candidate_info.get('Phone', 'N/A')}</p>
                                <p style="font-size: 1.1rem; margin: 1rem 0;"><strong>💼</strong> {candidate_info.get('Total_Experience', 'Not specified')}</p>
                                <p style="font-size: 1.1rem; margin: 1rem 0;"><strong>🎯</strong> {rec_emoji} {candidate.get('recommendation', 'CONSIDER')}</p>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            # Add expandable section for detailed analysis
                            with st.expander(f"📊 Detailed AI Analysis - {candidate_info.get('Name', 'Unknown')}", expanded=False):
                                col1, col2 = st.columns(2)
                                
                                with col1:
                                    st.markdown("#### 💡 AI Assessment")
                                    st.write(candidate.get('explanation', 'No explanation provided'))
                                    
                                    st.markdown("#### ✅ Key Strengths")
                                    strengths = candidate.get('key_strengths', ['Not specified'])
                                    for strength in strengths:
                                        st.write(f"• {strength}")
                                    
                                    st.markdown("#### 📈 Development Areas")
                                    gaps = candidate.get('key_gaps', ['None identified'])
                                    for gap in gaps:
                                        st.write(f"• {gap}")
                                
                                with col2:
                                    st.markdown("#### 📊 Detailed Score Breakdown")
                                    
                                    # Create score metrics with progress bars
                                    exp_score = match_details.get('experience', {}).get('score', 0)
                                    skills_score = match_details.get('skills', {}).get('score', 0)
                                    edu_score = match_details.get('education', {}).get('score', 0)
                                    cert_score = match_details.get('certifications', {}).get('score', 0)
                                    
                                    # Experience
                                    st.metric("🎯 Experience", f"{exp_score:.1f}%")
                                    st.progress(min(exp_score/100, 1.0))
                                    
                                    # Skills
                                    st.metric("🛠️ Skills", f"{skills_score:.1f}%")
                                    st.progress(min(skills_score/100, 1.0))
                                    
                                    # Education
                                    st.metric("🎓 Education", f"{edu_score:.1f}%")
                                    st.progress(min(edu_score/100, 1.0))
                                    
                                    # Certifications
                                    st.metric("📜 Certifications", f"{cert_score:.1f}%")
                                    st.progress(min(cert_score/100, 1.0))
                                    
                                    # Skills matching details
                                    matching_skills = match_details.get('skills', {}).get('matching_skills', [])
                                    if matching_skills:
                                        st.markdown("#### ✅ Matching Skills")
                                        skills_text = ", ".join(matching_skills[:8])  # Show up to 8 skills
                                        if len(matching_skills) > 8:
                                            skills_text += f" (+{len(matching_skills)-8} more)"
                                        st.info(skills_text)
                                    
                                    missing_skills = match_details.get('skills', {}).get('missing_skills', [])
                                    if missing_skills:
                                        st.markdown("#### ❌ Missing Skills")
                                        missing_text = ", ".join(missing_skills[:5])  # Show up to 5 missing skills
                                        if len(missing_skills) > 5:
                                            missing_text += f" (+{len(missing_skills)-5} more)"
                                        st.warning(missing_text)
                                    
                                    # Show transferable skills if available
                                    transferable_skills = match_details.get('skills', {}).get('transferable_skills', [])
                                    if transferable_skills:
                                        st.markdown("#### 🔄 Transferable Skills")
                                        trans_text = ", ".join(transferable_skills[:6])
                                        if len(transferable_skills) > 6:
                                            trans_text += f" (+{len(transferable_skills)-6} more)"
                                        st.success(trans_text)
                        
                        except Exception as e:
                            logger.error(f"Error displaying candidate {idx}: {e}")
                            st.error(f"Error displaying candidate {idx+1}")
                    
                    st.markdown("---")
                    
                    # Create summary table
                    st.markdown("### 📋 Complete Candidate Rankings")
                    
                    try:
                        summary_data = []
                        for idx, candidate in enumerate(st.session_state.top_candidates):
                            try:
                                candidate_info = candidate['candidate_data']
                                
                                # Get skills summary safely
                                skills = candidate_info.get("Skills", [])
                                if isinstance(skills, list):
                                    skills_summary = ", ".join(skills[:5])  # Show first 5 skills
                                    if len(skills) > 5:
                                        skills_summary += f" (+{len(skills)-5} more)"
                                else:
                                    skills_summary = str(skills)[:100] + "..." if len(str(skills)) > 100 else str(skills)
                                
                                summary_data.append({
                                    "Rank": f"#{idx + 1}",
                                    "Name": candidate_info.get("Name", "Unknown"),
                                    "Score": f"{candidate.get('overall_score', 0):.1f}%",
                                    "Recommendation": candidate.get('recommendation', 'CONSIDER'),
                                    "Experience": candidate_info.get("Total_Experience", "Not specified"),
                                    "Key Skills": skills_summary,
                                    "Email": candidate_info.get("Email", ""),
                                    "Phone": candidate_info.get("Phone", ""),
                                    "Resume File": candidate_info.get("Source_File", "")
                                })
                            except Exception as e:
                                logger.warning(f"Error processing candidate {idx} for summary: {e}")
                                continue
                        
                        summary_df = pd.DataFrame(summary_data)
                        
                        # Display with enhanced styling
                        st.dataframe(
                            summary_df, 
                            use_container_width=True,
                            height=400
                        )
                        
                        # Download section
                        st.markdown("---")
                        st.markdown("### 📥 Export Reports")
                        
                        col1, col2, col3, col4 = st.columns(4)
                        
                        with col1:
                            try:
                                   csv_data = summary_df.to_csv(index=False).encode('utf-8')
                                   st.download_button(
                                       "📊 All Candidates CSV",
                                       data=csv_data,
                                       file_name=f"all_candidates_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                       mime="text/csv",
                                       use_container_width=True
                                   )
                            except Exception as e:
                                   logger.error(f"Error creating CSV: {e}")
                                   st.error("Error creating CSV")
                       
                        with col2:
                            try:
                                   top_10_csv = summary_df.head(10).to_csv(index=False).encode('utf-8')
                                   st.download_button(
                                       "🏆 Top 10 CSV",
                                       data=top_10_csv,
                                       file_name=f"top_10_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                       mime="text/csv",
                                       use_container_width=True
                                   )
                            except Exception as e:
                                logger.error(f"Error creating top 10 CSV: {e}")
                                st.error("Error creating CSV")
                       
                        with col3:
                            try:
                                top_5_csv = summary_df.head(5).to_csv(index=False).encode('utf-8')
                                st.download_button(
                                    "🥇 Top 5 CSV",
                                    data=top_5_csv,
                                    file_name=f"top_5_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                    mime="text/csv",
                                    use_container_width=True
                                )
                            except Exception as e:
                                logger.error(f"Error creating top 5 CSV: {e}")
                                st.error("Error creating CSV")
                       
                        with col4:
                            try:
                                job_desc = st.session_state.get('current_job_description', "Job Description Not Available")
                                excel_data = create_excel_report(st.session_state.top_candidates, job_desc)
                                
                                if excel_data:
                                    st.download_button(
                                        "📈 Complete Report",
                                        data=excel_data,
                                        file_name=f"complete_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                        use_container_width=True
                                    )
                                else:
                                    st.error("Error creating Excel report")
                            except Exception as e:
                                logger.error(f"Error creating Excel report: {e}")
                                st.error("Error creating Excel report")
                   
                    except Exception as e:
                        logger.error(f"Error creating summary table: {e}")
                        st.error("Error creating candidate summary")
               
                else:
                   # Enhanced empty state
                   st.markdown("""
                   <div style="text-align: center; padding: 4rem 2rem; background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%); border-radius: 20px; margin: 2rem 0;">
                       <h3 style="color: #667eea; margin-bottom: 1rem;">🎯 No Candidates Analyzed Yet</h3>
                       <p style="color: #6c757d; font-size: 1.1rem; margin-bottom: 2rem;">Upload resumes and process them in the 'Upload & Process' tab to see AI-powered candidate rankings here.</p>
                       <div style="font-size: 4rem; margin: 2rem 0;">🤖</div>
                       <p style="color: #495057; font-style: italic;">AI-powered resume screening awaits your input!</p>
                   </div>
                   """, unsafe_allow_html=True)
           
            except Exception as e:
                logger.error(f"Error in shortlisted candidates tab: {e}")
                st.error(f"❌ Error displaying shortlisted candidates: {e}")
                st.info("💡 The application will continue running. Please try refreshing or contact support.")

    except Exception as e:
        logger.error(f"Critical error in main function: {e}")
        st.error(f"❌ Application error: {e}")
        st.info("💡 The application encountered an error but will continue running. Please try refreshing the page.")

if __name__ == "__main__":
   try:
       main()
   except Exception as e:
       logger.error(f"Critical startup error: {e}")
       st.error("❌ Failed to start the application. Please check your configuration and try again.")
       st.info("💡 Check that all required dependencies are installed and API keys are configured.")
